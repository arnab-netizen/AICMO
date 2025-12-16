"""DOCX validator."""

import os
from typing import Dict, Any
from .base import BaseValidator


class DOCXValidator(BaseValidator):
    """Validates DOCX documents."""
    
    def validate_structure(self, filepath: str, contract: Dict) -> Dict[str, Any]:
        """Validate DOCX structure."""
        result = {
            'valid': True,
            'errors': [],
            'warnings': []
        }
        
        if not os.path.exists(filepath):
            result['valid'] = False
            result['errors'].append(f"DOCX file not found: {filepath}")
            return result
        
        try:
            from docx import Document
            
            try:
                doc = Document(filepath)
                
                # Check paragraph count
                paragraph_count = len(doc.paragraphs)
                if paragraph_count == 0:
                    result['valid'] = False
                    result['errors'].append("DOCX has no paragraphs")
                
                # Check for content
                has_content = any(p.text.strip() for p in doc.paragraphs)
                if not has_content:
                    result['valid'] = False
                    result['errors'].append("DOCX has no text content")
                
                result['paragraph_count'] = paragraph_count
                result['has_content'] = has_content
                
            except Exception as e:
                result['valid'] = False
                result['errors'].append(f"DOCX read error: {str(e)}")
                
        except ImportError:
            result['warnings'].append("python-docx not installed, skipping DOCX validation")
        
        return result
    
    def validate_safety(self, filepath: str, contract: Dict) -> Dict[str, Any]:
        """Validate DOCX safety."""
        result = {
            'valid': True,
            'errors': [],
            'warnings': []
        }
        
        # Extract text
        text = self.extract_text(filepath)
        
        if not text:
            result['warnings'].append("No text extracted from DOCX")
            return result
        
        # Check placeholders
        placeholder_valid, placeholder_issues = self.validate_placeholders(text, contract)
        if not placeholder_valid:
            result['valid'] = False
            result['errors'].extend(placeholder_issues)
        
        # Check required strings
        required_valid, required_issues = self.validate_required_strings(text, contract)
        if not required_valid:
            result['valid'] = False
            result['errors'].extend(required_issues)
        
        # Check file size
        size_valid, size_issues = self.validate_file_size(filepath, contract)
        if not size_valid:
            result['valid'] = False
            result['errors'].extend(size_issues)
        
        result['no_placeholders'] = placeholder_valid
        result['required_strings_present'] = required_valid
        
        return result
    
    def extract_text(self, filepath: str) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document
            
            doc = Document(filepath)
            text_parts = [p.text for p in doc.paragraphs]
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except (ImportError, Exception):
            return ""
